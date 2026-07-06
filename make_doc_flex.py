from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
doc  = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def para(text): return doc.add_paragraph(text)

def mono(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.4)
    return p

def bullet(text, indent=0.4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.add_run(text); return p

def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]; cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, '2E74B5')
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        for ci, cell_text in enumerate(row):
            c = tr.cells[ci]; c.text = ''
            p = c.paragraphs[0]
            parts = cell_text.split('`')
            for idx, part in enumerate(parts):
                if not part: continue
                r = p.add_run(part)
                if idx % 2 == 1:
                    r.font.name = 'Courier New'; r.font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return t

def add_figure(img_path, width_inches=5.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].italic = True

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('FLEX lox2272/loxP Inversion — NGS Analysis Pipeline', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
doc.add_paragraph()

# ── Background ────────────────────────────────────────────────────────────────
h1('Background')
para(
    'The FLEX (Flip-Excision) construct carries a 2xYFP reporter cassette in antisense '
    'orientation, flanked by two pairs of incompatible lox sites in inverted orientation: '
    'loxP and lox2272. Cre recombinase inverts (flips) the cassette into sense orientation '
    'by recombining either the loxP pair or the lox2272 pair. Because loxP and lox2272 are '
    'incompatible (different 8 bp spacers: loxP = AGCATACAT; lox2272 = AGGATACTTT), they '
    'cannot cross-recombine — each pair acts independently.'
)
doc.add_paragraph()

# ── Construct Architecture ─────────────────────────────────────────────────────
h1('Construct Architecture and Molecular States')

h2('Construct layout')
para(
    'The pFLEX construct contains four lox sites arranged as two nested inverted pairs. '
    'The 2xYFP reporter cassette sits in antisense orientation between them.'
)
doc.add_paragraph()
mono('5\'--[lox2272->]A--L1--[loxP->]B--988bp antisense 2xYFP--[<-lox2272]C--L2--[<-loxP]D--3\'')
doc.add_paragraph()
para(
    'L1 and L2 are 50 bp linkers. loxP and lox2272 cannot cross-recombine (different spacer '
    'sequences), so each pair acts independently. In the starting state, both pairs are '
    'inverted — Cre can only invert, not excise.'
)
doc.add_paragraph()

h2('How Cre produces two distinct final products')
para(
    'Each recombination pathway requires two sequential Cre events to reach a permanently '
    'locked state. The two paths excise different linker regions, leaving different junction '
    'sequences — this is why two reference amplicons exist.'
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Path 1 — lox2272-flip then loxP excision').bold = True
mono('Step 1 (A-C inversion):  A[lox2272->] -- SENSE YFP -- B\'[<-loxP] -- L1_RC -- C[<-lox2272] -- L2 -- D[<-loxP]')
mono('Step 2 (B\'-D excision): A[lox2272->] -- SENSE YFP -- B\'[<-loxP]    <-- LOCKED')
para('L1 stays on the linear molecule. Junction retains the 8 bp segment CGCTGAC. => 584 bp amplicon.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Path 2 — loxP-flip then lox2272 excision').bold = True
mono('Step 1 (B-D inversion):  A[lox2272->] -- L1 -- B[loxP->] -- L2_RC -- C\'[lox2272->] -- SENSE YFP -- D[<-loxP]')
mono('Step 2 (A-C\' excision): A[lox2272->] -- SENSE YFP -- D[<-loxP]      <-- LOCKED')
para('L1 is excised into the circle. Junction lacks CGCTGAC. => 579 bp amplicon (5 bp shorter).')
doc.add_paragraph()

add_table(
    ['Product', 'Junction', 'Amplicon', 'Detected by'],
    [
        ['Unflipped',    'antisense 2xYFP intact',          '~928 bp', 'Unflip PCR'],
        ['Locked Path 1 (lox2272-flip)', 'retains CGCTGAC', '~584 bp', 'Flip PCR'],
        ['Locked Path 2 (loxP-flip)',    'lacks CGCTGAC',   '~579 bp', 'Flip PCR'],
    ],
    col_widths=[2.2, 2.0, 0.9, 1.2]
)
doc.add_paragraph()

h2('Why two separate PCRs are required')
para(
    'The post-anchor junction sequence is identical in unflipped and lox2272-flip reads — '
    'both have the same upstream flanking of site A. A single probe cannot distinguish them. '
    'The two-PCR design solves this at the amplification step: REV_UNFLIP only extends on '
    'unflipped molecules; REV_FLIP only extends on flipped molecules (it binds within the '
    'sense YFP, which is absent in the unflipped state). Flip reads are then classified into '
    'lox2272-flip vs loxP-flip by k-mer scoring of the 60 bp post-anchor window, exploiting '
    'the presence or absence of CGCTGAC.'
)
doc.add_paragraph()

# ── Conditions ────────────────────────────────────────────────────────────────
h1('Experimental Conditions')
add_table(
    ['Code', 'Condition', 'Cre delivery', 'n (bio reps)'],
    [
        ['c51',    'Negative control',          'None',       '2'],
        ['c51c39', 'Positive control',          'Agrobacterium', '2'],
        ['c30c41', 'PVC (normal)',               'PVC — empty vector', '2'],
        ['pFLS2',  'PVC pFLS2 overexpression',  'PVC — pFLS2::Cre',   '2 (reps 1,3)'],
        ['pUBQ10', 'PVC pUBQ10 overexpression', 'PVC — pUBQ10::Cre',  '2 (reps 2,3)'],
    ],
    col_widths=[0.8, 2.2, 2.2, 1.3]
)
doc.add_paragraph()

# ── Primers ───────────────────────────────────────────────────────────────────
h1('Primers and Key Sequences')

h2('PCR Primers')
add_table(
    ['Primer', "Sequence (5' to 3')", 'Used in', 'Length'],
    [
        ['Forward (shared)', '`CCACTGACGTAAGGGATGACGCACAATC`', 'Both reactions', '28 bp'],
        ['Reverse — unflip', '`GGGCGACACCCTGGTGAACCG`',       'Unflip PCR',     '21 bp'],
        ['Reverse — flip',   '`GCCCTCGAACTTCACCTCGGCG`',      'Flip PCR',       '22 bp'],
    ],
    col_widths=[1.5, 2.8, 1.5, 0.7]
)
doc.add_paragraph()

h2('Anchor Sequence')
para(
    'A 58 bp anchor sequence immediately downstream of the FWD primer is shared by all three '
    'molecular states (unflipped, loxP-flip, lox2272-flip). It is used to locate the read '
    'orientation and to define the start of the classification window.'
)
mono('Anchor (58 bp):')
mono('CCACTATCCTTCGCAAGACCCTTCCTCTATATAAGGAAGTTCATTTCATTTGGAGAGG')
doc.add_paragraph()

h2('Junction Probe Sequences')
para(
    'Classification uses a two-step motif + edit-distance strategy. '
    'lox2272-flip is detected by searching for the 7-mer CGCTGAC (≤ 1 mismatch) in '
    'a ±5 bp window around the post-anchor junction start. '
    'If the motif is absent, loxP-flip is called when the Levenshtein edit distance '
    'between the 30 bp post-anchor head and the loxP reference is ≤ 3. '
    'Edit distance (allowing substitutions and indels) correctly handles reads where '
    'a small indel at the junction start would cause Hamming distance to be spuriously '
    'inflated. Reads failing both tests are unclassified.'
)
add_table(
    ['State', '30 bp post-anchor head', 'Key motif', 'Amplicon'],
    [
        ['lox2272-flip', '`ACACGCTGACGGCCGCGTATTTTTACAACA`', '`CGCTGAC` (pos 3–9)', '584 bp'],
        ['loxP-flip',    '`ACAACGGCCGCGTATTTTTACAACAATTAC`', 'absent',              '579 bp'],
    ],
    col_widths=[1.2, 2.7, 1.4, 0.7]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Distinguishing feature: ').bold = True
p.add_run(
    'lox2272-flip retains the 7-mer `CGCTGAC` immediately before `GGCCGCG`; '
    'loxP-flip lacks this segment entirely. The motif search in a ±5 bp window '
    'correctly handles lox2272 reads whose anchor was detected ~5 bp late. '
    'Edit distance (rather than Hamming) is used for loxP confirmation because '
    'Nanopore indels at the junction start cause a frameshift that inflates Hamming '
    'distance — two confirmed loxP reads had Hamming = 37/30 but edit distance = 3.'
)
doc.add_paragraph()

# ── Pipeline ──────────────────────────────────────────────────────────────────
h1('Pipeline Overview')

p = doc.add_paragraph()
r = p.add_run('Step 1: Size filter')
r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
bullet('Unflip PCR reads: 800–1,050 bp (expected ~928 bp)')
bullet('Flip PCR reads: 480–680 bp (expected 579 or 584 bp)')
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Step 2: FWD primer detection and strand normalisation')
r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
bullet('Hamming search (≤ 4 mismatches) on both the read and its reverse complement')
bullet('The strand on which FWD falls earliest defines canonical orientation')
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Step 3: Anchor detection')
r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
bullet('Hamming search (≤ 8 mismatches) for the 58 bp anchor in the window immediately after FWD')
bullet('Unflip reads that pass anchor detection are counted as confirmed unflipped molecules')
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Step 4: Junction classification — motif + edit distance (flip reads only)')
r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
bullet('Step 4a (lox2272): search for CGCTGAC (≤ 1 mismatch) in a ±5 bp window around '
       'the post-anchor start. Found → lox2272-flip.')
bullet('Step 4b (loxP): if motif absent, compute Levenshtein edit distance between the '
       '30 bp post-anchor head and the loxP reference. ≤ 3 edits → loxP-flip; else unclassified.')
bullet('Edit distance (not Hamming) is critical for loxP: a single indel at the junction '
       'start causes Hamming distance to cascade to ~37/30 while edit distance remains ≤ 3.')
bullet(
    'Why flip reads only: the post-anchor junction is identical in unflipped and '
    'lox2272-flip molecules. Unflipped reads are already separated at the PCR level '
    '(REV_UNFLIP only amplifies unflipped molecules; REV_FLIP only amplifies flipped ones).'
)
doc.add_paragraph()

# ── Results ───────────────────────────────────────────────────────────────────
h1('Results')

h2('Table 1 — Per-sample read counts and flip classification')
para('Total reads (unflip) and Total reads (flip) are raw read counts from each PCR reaction. '
     'Unflip (N) = reads passing size filter + FWD primer + anchor in the unflip reaction. '
     'lox2272-flip: CGCTGAC motif (≤ 1 mm, ±5 bp window). '
     'loxP-flip: edit distance ≤ 3 on 30 bp post-anchor head. See Pipeline for details.')
add_table(
    ['Condition', 'Rep',
     'Total\n(unflip)', 'Total\n(flip)',
     'Unflip (N)', 'loxP-flip', 'lox2272-flip', 'Total flip', '%loxP', '%lox2272'],
    [
        ['Neg',          '1', '5,470', '5,482', '4,484',  '0',    '1,566',  '1,566',  '0.00%', '100.00%'],
        ['Neg',          '2', '5,466', '3,988', '4,706',  '0',       '92',     '92',  '0.00%', '100.00%'],
        ['Agro (+ctrl)', '1', '5,473', '5,467', '4,729',  '1',    '4,573',  '4,574',  '0.02%',  '99.98%'],
        ['Agro (+ctrl)', '2', '5,457', '5,463', '4,692',  '0',    '4,352',  '4,352',  '0.00%', '100.00%'],
        ['PVC (normal)', '1', '5,461', '5,463', '4,548',  '0',    '3,441',  '3,441',  '0.00%', '100.00%'],
        ['PVC (normal)', '2', '5,472', '5,467', '4,612',  '3',    '2,601',  '2,604',  '0.12%',  '99.88%'],
        ['PVC pFLS2',    '1', '5,493', '5,478', '4,551',  '0',    '4,291',  '4,291',  '0.00%', '100.00%'],
        ['PVC pFLS2',    '3', '5,479', '5,473', '4,636',  '0',    '4,364',  '4,364',  '0.00%', '100.00%'],
        ['PVC pUBQ10',   '2', '5,448', '5,470', '4,599',  '0',    '4,143',  '4,143',  '0.00%', '100.00%'],
        ['PVC pUBQ10',   '3', '5,462', '5,478', '4,601',  '0',    '4,347',  '4,347',  '0.00%', '100.00%'],
    ],
    col_widths=[1.0, 0.32, 0.65, 0.62, 0.65, 0.6, 0.8, 0.68, 0.6, 0.7]
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Note on Neg rep1: ').bold = True
p.add_run(
    'The negative control replicate 1 shows an anomalously high lox2272-flip count '
    '(1,566 reads) compared with replicate 2 (92 reads). This discrepancy likely reflects '
    'a sample-level issue (e.g. low-level Cre contamination or sample swap) and should be '
    'confirmed by re-sequencing or re-planting. Neg rep1 is included for transparency but '
    'is treated as an outlier.'
)
doc.add_paragraph()

h2('Table 2 — Junction alignment to lox2272-flip reference (non-Neg conditions)')
add_table(
    ['Condition', 'lox2272-flip reads (pooled)', 'Avg identity to reference', 'Min per-position identity'],
    [
        ['Agro (+ctrl)',  '8,925', '99.7%', '97.9% (pos 60)'],
        ['PVC (normal)',  '6,042', '99.7%', '97.9% (pos 60)'],
        ['PVC pFLS2',    '8,655', '99.7%', '97.4% (pos 60)'],
        ['PVC pUBQ10',   '8,490', '99.8%', '97.5% (pos 60)'],
    ],

    col_widths=[1.5, 1.8, 1.7, 2.0]
)
doc.add_paragraph()
para(
    'All four conditions show near-perfect junction sequence fidelity. The slight drop at '
    'position 60 (3\' end of the classification window) is consistent with Nanopore terminal '
    'errors and is not indicative of a biological variant. All 60 positions exceed 97% '
    'per-read identity, confirming the correct lox2272-mediated inversion junction.'
)
doc.add_paragraph()

h2('Figure 1 — Total flipped read counts per condition')
add_figure(os.path.join(BASE, 'Figure 1 — Total flipped read counts per condition.png'), width_inches=4.8,
           caption='Figure 1. Total flipped reads (lox2272-flip + loxP-flip) per condition '
                   '(individual replicates shown, bar = mean). '
                   'Neg rep1* is anomalously high and treated as an outlier (see text).')
doc.add_paragraph()

h2('Figure 2 — Per-position junction identity to reference')
add_figure(os.path.join(BASE, 'Figure 2 — Per-position junction identity to reference.png'), width_inches=6.0,
           caption='Figure 2. Per-position base identity (%) of lox2272-flip reads to the '
                   '60 bp reference junction, pooled across both replicates per condition. '
                   'All four conditions overlap closely (avg 99.7–99.8%).')
doc.add_paragraph()

h2('Figure 3 — Flip read composition per replicate')
add_figure(os.path.join(BASE, 'Figure 3 — Flip read composition per replicate.png'), width_inches=5.5,
           caption='Figure 3. Proportion of flip reads classified as lox2272-flip (purple) '
                   'vs loxP-flip (orange) per replicate (non-negative conditions only). '
                   'lox2272-flip accounts for ≥99% of all classified flip events.')
doc.add_paragraph()

h2('Figure 4 — lox2272-flip percentage per condition')
add_figure(os.path.join(BASE, 'Figure 4 — lox2272-flip percentage per condition.png'), width_inches=4.8,
           caption='Figure 4. lox2272-flip as a percentage of all classified flip reads '
                   '(individual replicates, bar = mean; y-axis 0–100%). '
                   'All conditions are consistently >99%.')
doc.add_paragraph()

h2('Figure 5 — loxP-flip percentage per condition')
add_figure(os.path.join(BASE, 'Figure 5 — loxP-flip percentage per condition.png'), width_inches=4.8,
           caption='Figure 5. loxP-flip as a percentage of all classified flip reads '
                   '(individual replicates, bar = mean; y-axis starts at 0). '
                   'loxP-flip accounts for 0.0–0.1% of flip events across all conditions.')
doc.add_paragraph()

h2('Figure 6 — lox2272-flip junction alignment')
add_figure(os.path.join(BASE, 'Figure 6 — lox2272-flip junction alignment.png'), width_inches=6.5,
           caption='Figure 6. Nucleotide alignment of the most common lox2272-flip junction sequence '
                   'per condition (most frequent read shown). Last 20 bp of the 35S anchor through '
                   'the first 20 bp of 2×YFP (184 bp total). Annotation bars: 35S promoter, '
                   'lox2272 junction (spacer), TMV UTR, 2×YFP. All conditions share an identical '
                   'consensus junction sequence. A=#8dd3c7, T=#fb8072, G=#ffffb3, C=#bebada.')
doc.add_paragraph()

h2('Figure 7 — Unique lox2272-flip junction sequences per sample (≤5 mismatches, n ≥ 5)')
para(
    'Each panel shows all unique junction sequences (184 bp window) with ≤5 mismatches to the '
    'reference and ≥5 supporting reads, ranked by read count. Mismatches from reference are '
    'highlighted with a red border. One panel per replicate; sequences are predominantly '
    'identical to reference, confirming high-fidelity Cre-mediated inversion at the lox2272 site.'
)
doc.add_paragraph()

FIG7_SAMPLES = [
    ('Neg',          'rep 1'),
    ('Neg',          'rep 2'),
    ('Agro (+ctrl)', 'rep 1'),
    ('Agro (+ctrl)', 'rep 2'),
    ('PVC (normal)', 'rep 1'),
    ('PVC (normal)', 'rep 2'),
    ('PVC pFLS2',    'rep 1'),
    ('PVC pFLS2',    'rep 2'),
    ('PVC pUBQ10',   'rep 1'),
    ('PVC pUBQ10',   'rep 2'),
]
for condition, rep in FIG7_SAMPLES:
    fname = f'Figure 7 — {condition} {rep} unique junction sequences.png'
    fpath = os.path.join(BASE, fname)
    if os.path.exists(fpath):
        add_figure(fpath, width_inches=6.5,
                   caption=f'Figure 7 ({condition} {rep}). Unique lox2272-flip junction sequences '
                           f'with ≤5 mismatches and ≥5 reads.')
        doc.add_paragraph()
    else:
        para(f'[Figure 7 — {condition} {rep}: file not found]')
        doc.add_paragraph()

# ── Key findings ──────────────────────────────────────────────────────────────
h1('Key Findings')

bullet('lox2272-mediated inversion dominates across all conditions and replicates '
       '(99.9–100.0% of classified flip reads), confirming strong preference for '
       'lox2272-site recombination in this FLEX construct.')
bullet('loxP-flip is extremely rare: 0–3 reads per sample (0.0–0.1%), detected only in '
       'Agro (+ctrl) rep1 (1 read) and PVC (normal) rep2 (3 reads). These counts are '
       'at the noise floor and do not differ meaningfully between conditions.')
bullet('Agro (+ctrl), PVC pFLS2, and PVC pUBQ10 show high and consistent flip read counts '
       '(~4,100–4,580 per replicate), similar to the Agrobacterium positive control.')
bullet('PVC (normal) shows intermediate counts (~2,600–3,440), suggesting lower Cre '
       'delivery efficiency without the FLS2 or UBQ10 promoter-driven boost.')
bullet('Neg rep2 shows the expected near-zero flip signal (92 reads), confirming the '
       'pipeline correctly identifies unflipped molecules. Neg rep1 is anomalous and '
       'should be re-examined.')
bullet('Junction sequences in all four non-Neg conditions match the lox2272-flip '
       'reference with average 99.7% per-position identity, confirming correct '
       'Cre-mediated inversion at the expected lox2272 site.')
doc.add_paragraph()

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, 'FLEX_NGS_analysis_summary_v14.docx')
doc.save(out)
print('Saved:', out)
